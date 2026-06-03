import json

from docx import Document
from typer.testing import CliRunner

from heitang_kb_forge.cli import app
from heitang_kb_forge.parsers.docx_parser import parse_docx

STANDARD_PACKAGE_FILES = {
    "chunks.jsonl",
    "cards.jsonl",
    "qa_pairs.jsonl",
    "glossary.jsonl",
    "manifest.json",
    "ingest_report.md",
    "quality_report.json",
}


def test_docx_parser_extracts_text_from_text_based_docx(tmp_path):
    sample_docx = tmp_path / "sample_text.docx"
    _write_minimal_text_docx(sample_docx)

    text = parse_docx(sample_docx)

    assert isinstance(text, str)
    assert "KB Forge DOCX Fixture" in text
    assert "text-based DOCX parsing" in text


def test_build_processes_text_based_docx(tmp_path):
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    input_dir.mkdir()
    _write_minimal_text_docx(input_dir / "sample_text.docx")

    result = CliRunner().invoke(
        app,
        [
            "build",
            "--input",
            str(input_dir),
            "--output",
            str(output_dir),
            "--domain",
            "education",
            "--mode",
            "teaching",
        ],
    )

    assert result.exit_code == 0, result.output
    assert (output_dir / "chunks.jsonl").exists()
    assert (output_dir / "manifest.json").exists()
    assert (output_dir / "ingest_report.md").exists()

    chunk_lines = (output_dir / "chunks.jsonl").read_text(encoding="utf-8").splitlines()
    assert len(chunk_lines) >= 1
    first_chunk = json.loads(chunk_lines[0])
    assert "KB Forge DOCX Fixture" in first_chunk["text"]
    assert "text-based DOCX parsing" in first_chunk["text"]


def test_docx_parser_extracts_single_table(tmp_path):
    sample_docx = tmp_path / "table.docx"
    document = Document()
    table = document.add_table(rows=2, cols=3)
    _set_row(table.rows[0], ["书名", "作者", "定价"])
    _set_row(table.rows[1], ["产品经理入门", "张三", "59"])
    document.save(sample_docx)

    text = parse_docx(sample_docx)

    assert "Table 1. Row 2. 书名: 产品经理入门. 作者: 张三. 定价: 59." in text


def test_docx_parser_extracts_multiple_tables_in_order(tmp_path):
    sample_docx = tmp_path / "tables.docx"
    document = Document()
    first = document.add_table(rows=2, cols=1)
    _set_row(first.rows[0], ["Name"])
    _set_row(first.rows[1], ["First Table Fixture"])
    second = document.add_table(rows=2, cols=1)
    _set_row(second.rows[0], ["Name"])
    _set_row(second.rows[1], ["Second Table Fixture"])
    document.save(sample_docx)

    text = parse_docx(sample_docx)

    assert "Table 1. Row 2. Name: First Table Fixture." in text
    assert "Table 2. Row 2. Name: Second Table Fixture." in text
    assert text.index("Table 1. Row 2") < text.index("Table 2. Row 2")


def test_docx_parser_handles_empty_and_duplicate_headers_and_skips_empty_rows(tmp_path):
    sample_docx = tmp_path / "headers.docx"
    document = Document()
    table = document.add_table(rows=4, cols=3)
    _set_row(table.rows[0], ["Name", "", "Name"])
    _set_row(table.rows[1], ["Course", "99", "Duplicate"])
    _set_row(table.rows[2], ["", "", ""])
    _set_row(table.rows[3], ["Another", "", "Value"])
    document.save(sample_docx)

    text = parse_docx(sample_docx)

    assert "Table 1. Row 2. Name: Course. Column B: 99. Name 2: Duplicate." in text
    assert "Table 1. Row 3" not in text
    assert "Table 1. Row 4. Name: Another. Name 2: Value." in text


def test_docx_parser_uses_column_names_when_first_table_row_is_empty(tmp_path):
    sample_docx = tmp_path / "no_header.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    _set_row(table.rows[0], ["", ""])
    _set_row(table.rows[1], ["Course", "99"])
    document.save(sample_docx)

    text = parse_docx(sample_docx)

    assert "Table 1. Row 2. Column A: Course. Column B: 99." in text


def test_docx_parser_returns_paragraphs_and_tables(tmp_path):
    sample_docx = tmp_path / "mixed.docx"
    document = Document()
    document.add_paragraph("Paragraph Fixture")
    table = document.add_table(rows=2, cols=1)
    _set_row(table.rows[0], ["Name"])
    _set_row(table.rows[1], ["Table Fixture"])
    document.save(sample_docx)

    text = parse_docx(sample_docx)

    assert "Paragraph Fixture" in text
    assert "Table 1. Row 2. Name: Table Fixture." in text


def test_build_processes_docx_table_and_writes_standard_package(tmp_path):
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    input_dir.mkdir()
    _write_table_docx(input_dir / "table.docx", "KB Forge DOCX Table Fixture")

    result = CliRunner().invoke(app, ["build", "--input", str(input_dir), "--output", str(output_dir)])

    assert result.exit_code == 0, result.output
    assert {path.name for path in output_dir.iterdir()} == STANDARD_PACKAGE_FILES
    assert (output_dir / "quality_report.json").exists()
    assert "KB Forge DOCX Table Fixture" in _read_chunk_text(output_dir)


def test_batch_processes_docx_table(tmp_path):
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    input_dir.mkdir()
    _write_table_docx(input_dir / "001_table.docx", "KB Forge DOCX Table Fixture")

    result = CliRunner().invoke(app, ["batch", "--input", str(input_dir), "--output", str(output_dir)])

    assert result.exit_code == 0, result.output
    assert {path.name for path in (output_dir / "001_table").iterdir()} == STANDARD_PACKAGE_FILES
    assert "KB Forge DOCX Table Fixture" in _read_chunk_text(output_dir / "001_table")


def test_merge_combines_docx_table_markdown_and_csv(tmp_path):
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    input_dir.mkdir()
    (input_dir / "001_notes.md").write_text("KB Forge Markdown Fixture", encoding="utf-8")
    (input_dir / "001_table.csv").write_text("Name\nKB Forge CSV Fixture\n", encoding="utf-8")
    _write_table_docx(input_dir / "001_docx_table.docx", "KB Forge DOCX Table Fixture")

    result = CliRunner().invoke(
        app,
        ["batch", "--input", str(input_dir), "--output", str(output_dir), "--merge-same-sequence"],
    )

    assert result.exit_code == 0, result.output
    assert {path.name for path in (output_dir / "001").iterdir()} == STANDARD_PACKAGE_FILES
    chunk_text = _read_chunk_text(output_dir / "001")
    assert "KB Forge Markdown Fixture" in chunk_text
    assert "KB Forge CSV Fixture" in chunk_text
    assert "KB Forge DOCX Table Fixture" in chunk_text


def _write_minimal_text_docx(path):
    document = Document()
    document.add_paragraph("KB Forge DOCX Fixture")
    document.add_paragraph("text-based DOCX parsing")
    document.save(path)


def _write_table_docx(path, value):
    document = Document()
    table = document.add_table(rows=2, cols=1)
    _set_row(table.rows[0], ["Name"])
    _set_row(table.rows[1], [value])
    document.save(path)


def _set_row(row, values):
    for cell, value in zip(row.cells, values):
        cell.text = value


def _read_chunk_text(output_dir):
    return "\n".join(
        json.loads(line)["text"] for line in (output_dir / "chunks.jsonl").read_text(encoding="utf-8").splitlines()
    )
