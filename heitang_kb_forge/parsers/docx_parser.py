from pathlib import Path
from typing import Iterable


def parse_docx(path: Path) -> str:
    try:
        from docx import Document

        document = Document(path)
    except Exception:
        return ""

    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_paragraphs = _format_tables(document.tables)
    return "\n\n".join(paragraphs + table_paragraphs)


def _format_tables(tables) -> list[str]:
    paragraphs: list[str] = []
    for table_index, table in enumerate(tables, start=1):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        paragraphs.extend(_format_table_rows(rows, table_index))
    return paragraphs


def _format_table_rows(rows: list[list[object]], table_index: int) -> list[str]:
    if not rows:
        return []

    width = max((len(row) for row in rows), default=0)
    if width == 0:
        return []

    first_row = _pad_row(rows[0], width)
    if _row_has_values(first_row):
        headers = _normalize_headers(first_row)
        data_rows = rows[1:]
        row_offset = 2
    else:
        headers = _normalize_headers([])
        data_rows = rows
        row_offset = 1

    headers = _pad_headers(headers, width)
    paragraphs: list[str] = []
    for row_index, row in enumerate(data_rows, start=row_offset):
        padded = _pad_row(row, width)
        if not _row_has_values(padded):
            continue
        fields = [
            f"{headers[column_index]}: {_stringify(value)}"
            for column_index, value in enumerate(padded)
            if _stringify(value)
        ]
        if fields:
            paragraphs.append(f"Table {table_index}. Row {row_index}. {'. '.join(fields)}.")

    return paragraphs


def _normalize_headers(row: Iterable[object]) -> list[str]:
    headers: list[str] = []
    seen: dict[str, int] = {}
    for index, value in enumerate(row):
        header = _stringify(value) or _column_name(index)
        count = seen.get(header, 0) + 1
        seen[header] = count
        headers.append(header if count == 1 else f"{header} {count}")
    return headers


def _pad_headers(headers: list[str], width: int) -> list[str]:
    return headers + [_column_name(index) for index in range(len(headers), width)]


def _pad_row(row: list[object], width: int) -> list[object]:
    return row + [""] * (width - len(row))


def _row_has_values(row: Iterable[object]) -> bool:
    return any(_stringify(value) for value in row)


def _stringify(value: object) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _column_name(index: int) -> str:
    name = ""
    number = index + 1
    while number:
        number, remainder = divmod(number - 1, 26)
        name = chr(65 + remainder) + name
    return f"Column {name}"
