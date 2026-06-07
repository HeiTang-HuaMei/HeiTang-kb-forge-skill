from __future__ import annotations

import re
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape

from heitang_kb_forge.document_generation.planner import plan_document_generation
from heitang_kb_forge.document_generation.reporter import (
    generated_file_report,
    generation_trace,
    quality_report,
    render_export_validation_report,
    render_generated_file_report,
)
from heitang_kb_forge.document_generation.templates import render_markdown
from heitang_kb_forge.document_generation.validators import validate_exports
from heitang_kb_forge.exporters.jsonl_exporter import write_json

DOCUMENT_GENERATION_OUTPUT_FILES = [
    "generated.md",
    "generated.docx",
    "generated.pdf",
    "generated.pptx",
    "generated_file_report.json",
    "generated_file_report.md",
    "document_generation_trace.json",
    "document_quality_report.json",
    "export_validation_report.json",
    "export_validation_report.md",
]
SUPPORTED_FORMATS = {"md", "docx", "pdf", "pptx"}


def generate_document_outputs(
    package: Path,
    output: Path,
    formats: list[str],
    template: str = "default",
    grounding_policy: str = "strict_grounded",
    title: str | None = None,
) -> dict:
    requested = _normalize_formats(formats)
    output.mkdir(parents=True, exist_ok=True)
    plan = plan_document_generation(package, template, grounding_policy, title)
    markdown = render_markdown(plan)
    generated = _write_formats(output, requested, markdown, plan.title)
    validation = validate_exports(generated)
    file_report = generated_file_report(generated)
    write_json(output / "export_validation_report.json", validation)
    (output / "export_validation_report.md").write_text(render_export_validation_report(validation), encoding="utf-8")
    write_json(output / "generated_file_report.json", file_report)
    (output / "generated_file_report.md").write_text(render_generated_file_report(file_report), encoding="utf-8")
    write_json(output / "document_generation_trace.json", generation_trace(plan, generated, validation))
    write_json(output / "document_quality_report.json", quality_report(plan, markdown, validation))
    return {
        "status": validation["status"],
        "formats": requested,
        "files": {fmt: str(path) for fmt, path in sorted(generated.items())},
        "review_required": plan.review_required,
        "warnings": plan.warnings,
    }


def _normalize_formats(formats: list[str]) -> list[str]:
    requested = [fmt.lower().lstrip(".") for fmt in formats]
    unknown = sorted(set(requested) - SUPPORTED_FORMATS)
    if unknown:
        raise ValueError(f"Unsupported document format(s): {', '.join(unknown)}")
    return requested or ["md"]


def _write_formats(output: Path, formats: list[str], markdown: str, title: str) -> dict[str, Path]:
    files: dict[str, Path] = {}
    if "md" in formats:
        path = output / "generated.md"
        path.write_text(markdown, encoding="utf-8")
        files["md"] = path
    if "docx" in formats:
        path = output / "generated.docx"
        _write_docx(path, markdown)
        files["docx"] = path
    if "pdf" in formats:
        path = output / "generated.pdf"
        _write_pdf(path, markdown)
        files["pdf"] = path
    if "pptx" in formats:
        path = output / "generated.pptx"
        _write_pptx(path, title, markdown)
        files["pptx"] = path
    return files


def _write_docx(path: Path, markdown: str) -> None:
    try:
        from docx import Document
    except ModuleNotFoundError:
        _write_minimal_docx(path, markdown)
        return
    document = Document()
    for line in markdown.splitlines():
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            document.add_paragraph(line)
    document.save(path)


def _write_minimal_docx(path: Path, markdown: str) -> None:
    body = "".join(f"<w:p><w:r><w:t>{escape(line)}</w:t></w:r></w:p>" for line in markdown.splitlines() if line.strip())
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"<w:body>{body}<w:sectPr/></w:body></w:document>"
    )
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", _docx_content_types())
        archive.writestr("_rels/.rels", _docx_rels())
        archive.writestr("word/document.xml", document_xml)


def _write_pdf(path: Path, markdown: str) -> None:
    lines = [_pdf_safe(line) for line in _plain_lines(markdown)[:42]]
    text_ops = ["BT", "/F1 11 Tf", "50 780 Td", "14 TL"]
    for line in lines:
        text_ops.append(f"({line}) Tj")
        text_ops.append("T*")
    text_ops.append("ET")
    stream = "\n".join(text_ops).encode("latin-1", errors="replace")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    _write_pdf_objects(path, objects)


def _write_pdf_objects(path: Path, objects: list[bytes]) -> None:
    payload = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(payload))
        payload.extend(f"{index} 0 obj\n".encode("ascii"))
        payload.extend(obj)
        payload.extend(b"\nendobj\n")
    xref_offset = len(payload)
    payload.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    payload.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        payload.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    payload.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(bytes(payload))


def _write_pptx(path: Path, title: str, markdown: str) -> None:
    bullet_text = "\n".join(_plain_lines(markdown)[1:7])
    slide = _pptx_slide_xml(title, bullet_text)
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", _pptx_content_types())
        archive.writestr("_rels/.rels", _pptx_root_rels())
        archive.writestr("ppt/presentation.xml", _pptx_presentation_xml())
        archive.writestr("ppt/_rels/presentation.xml.rels", _pptx_presentation_rels())
        archive.writestr("ppt/slides/slide1.xml", slide)


def _plain_lines(markdown: str) -> list[str]:
    lines = []
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        line = re.sub(r"^#{1,3}\s+", "", line)
        line = line.removeprefix("- ").replace("`", "")
        lines.append(line)
    return lines


def _pdf_safe(text: str) -> str:
    return text[:92].replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _docx_content_types() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        "</Types>"
    )


def _docx_rels() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
        "</Relationships>"
    )


def _pptx_content_types() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/ppt/presentation.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml"/>'
        '<Override PartName="/ppt/slides/slide1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slide+xml"/>'
        "</Types>"
    )


def _pptx_root_rels() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="ppt/presentation.xml"/>'
        "</Relationships>"
    )


def _pptx_presentation_xml() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<p:presentation xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<p:sldIdLst><p:sldId id="256" r:id="rId1"/></p:sldIdLst>'
        '<p:sldSz cx="9144000" cy="5143500" type="screen16x9"/>'
        "</p:presentation>"
    )


def _pptx_presentation_rels() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide1.xml"/>'
        "</Relationships>"
    )


def _pptx_slide_xml(title: str, body: str) -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        "<p:cSld><p:spTree>"
        '<p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr/>'
        f'{_pptx_text_box(2, "Title", 457200, 320000, 8229600, 700000, title)}'
        f'{_pptx_text_box(3, "Body", 457200, 1180000, 8229600, 3000000, body)}'
        "</p:spTree></p:cSld></p:sld>"
    )


def _pptx_text_box(shape_id: int, name: str, x: int, y: int, cx: int, cy: int, text: str) -> str:
    paragraphs = "".join(f"<a:p><a:r><a:t>{escape(line[:220])}</a:t></a:r></a:p>" for line in text.splitlines() if line.strip())
    return (
        f'<p:sp><p:nvSpPr><p:cNvPr id="{shape_id}" name="{name}"/><p:cNvSpPr txBox="1"/><p:nvPr/></p:nvSpPr>'
        f'<p:spPr><a:xfrm><a:off x="{x}" y="{y}"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></p:spPr>'
        f"<p:txBody><a:bodyPr/><a:lstStyle/>{paragraphs}</p:txBody></p:sp>"
    )
